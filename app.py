from flask import Flask, render_template, request, send_file, redirect, url_for, flash, session
from docx import Document
from datetime import date, datetime, timezone
import io
import json
import requests
from functools import wraps

# --- INICIALIZAÇÃO DO FIREBASE ---
import firebase_admin
from firebase_admin import credentials, firestore

# Certifique-se de que o arquivo 'firebase-credentials.json' está na mesma pasta
try:
    cred = credentials.Certificate("firebase-credentials.json")
    firebase_admin.initialize_app(cred)
    db = firestore.client()
except Exception as e:
    print(f"Erro ao inicializar o Firebase: {e}")
    db = None

# --- CONFIGURAÇÃO DO APP FLASK ---
app = Flask(__name__)
app.secret_key = 'Xablau' # Mantenha sua chave secreta

# IMPORTANTE: Chave da API Web do Firebase
FIREBASE_WEB_API_KEY = "AIzaSyDMhMm78EASXSSF8Y6uBwLvcenGkWkrf00"

# --- FUNÇÃO DE GERAÇÃO DO DOCX (Contrato Antigo - Mantida para referência) ---
def gerar_contrato_docx(dados):
    settings_ref = db.collection('configuracoes').document('empresa')
    settings_doc = settings_ref.get()
    empresa_info = settings_doc.to_dict() if settings_doc.exists else {}
    nome_empresa = empresa_info.get('nome_empresa', 'GL Technology – Tecnologia para Crescer')
    cnpj_empresa = empresa_info.get('cnpj', '[SEU CNPJ AINDA NÃO CONFIGURADO]')
    endereco_empresa = empresa_info.get('endereco', '[SEU ENDEREÇO AINDA NÃO CONFIGURADO]')

    doc = Document()
    doc.add_heading('CONTRATO DE PRESTAÇÃO DE SERVIÇOS DE TECNOLOGIA', level=1)
    doc.add_paragraph()
    p_contratante = doc.add_paragraph('Pelo presente instrumento, de um lado: ')
    p_contratante.add_run('CONTRATANTE: ').bold = True
    p_contratante.add_run(f'{dados["nome_cliente"]}, inscrito(a) no CPF/CNPJ sob o nº {dados["doc_cliente"]}, doravante denominado(a) simplesmente CONTRATANTE.')
    p_contratada = doc.add_paragraph('De outro lado: ')
    p_contratada.add_run('CONTRATADA: ').bold = True
    p_contratada.add_run(f'{nome_empresa}, pessoa jurídica de direito privado, inscrita no CNPJ sob o nº {cnpj_empresa}, com sede em {endereco_empresa}, doravante denominada simplesmente CONTRATADA.')
    doc.add_paragraph()
    doc.add_heading('CLÁUSULA 1ª - DO OBJETO', level=2)
    doc.add_paragraph(f'1.1. O objeto do presente contrato é a prestação de serviços de {dados["servico"]}.')
    doc.add_heading('CLÁUSULA 2ª - DO PRAZO', level=2)
    doc.add_paragraph(f'2.1. O serviço objeto deste contrato será executado no prazo de {dados["prazo_dias"]} dias úteis.')
    doc.add_heading('CLÁUSULA 3ª - DO VALOR E FORMA DE PAGAMENTO', level=2)
    doc.add_paragraph(f'3.1. Pela prestação dos serviços, a CONTRATANTE pagará à CONTRATADA o valor total de R$ {float(dados["valor_total"]):.2f}.')
    data_formatada = dados.get('data_criacao', date.today()).strftime("%d de %B de %Y")
    doc.add_paragraph(f'\n[Local de Assinatura], {data_formatada}.')
    doc.add_paragraph('\n\n______________________________________')
    doc.add_paragraph(f'{dados["nome_cliente"]}\nCONTRATANTE')
    doc.add_paragraph('\n\n______________________________________')
    doc.add_paragraph(f'{nome_empresa}\nCONTRATADA')
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- SISTEMA DE AUTENTICAÇÃO ---
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user' not in session:
            flash("Você precisa fazer login para acessar esta página.", "error")
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

@app.route("/login", methods=["GET", "POST"])
def login():
    if 'user' in session:
        return redirect(url_for('dashboard'))
    if request.method == "POST":
        email = request.form.get('email')
        password = request.form.get('password')
        rest_api_url = f"https://identitytoolkit.googleapis.com/v1/accounts:signInWithPassword?key={FIREBASE_WEB_API_KEY}"
        payload = json.dumps({"email": email, "password": password, "returnSecureToken": True})
        try:
            r = requests.post(rest_api_url, data=payload)
            r.raise_for_status()
            data = r.json()
            if 'idToken' in data:
                session['user'] = data['email']
                session['idToken'] = data['idToken']
                return redirect(url_for('dashboard'))
        except requests.exceptions.HTTPError:
            flash("Credenciais inválidas. Verifique seu e-mail e senha.", "error")
        except Exception as e:
            flash(f"Ocorreu um erro de conexão: {e}", "error")
    return render_template("login.html")

@app.route("/logout")
def logout():
    session.pop('user', None)
    session.pop('idToken', None)
    flash("Você saiu da sua conta.", "success")
    return redirect(url_for('login'))

# --- ROTAS PRINCIPAIS DA APLICAÇÃO ---

@app.route("/")
@login_required
def dashboard():
    now = datetime.now(timezone.utc)
    start_of_month = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
    
    total_os = 0
    valor_total_gerado = 0
    os_este_mes = 0
    valor_este_mes = 0

    docs_gerais = db.collection('ordens_servico').stream()
    for doc in docs_gerais:
        os_data = doc.to_dict()
        total_os += 1
        valor_os = os_data.get('orcamento_servicos', 0) + os_data.get('orcamento_produtos', 0)
        valor_total_gerado += valor_os
        
        data_abertura = os_data.get('data_abertura')
        if data_abertura and data_abertura >= start_of_month:
            os_este_mes += 1
            valor_este_mes += valor_os

    recentes_ref = db.collection('ordens_servico').order_by('data_abertura', direction=firestore.Query.DESCENDING).limit(5).stream()
    os_recentes = []
    for os in recentes_ref:
        dados = os.to_dict()
        dados['id'] = os.id
        dados['data_abertura_fmt'] = dados['data_abertura'].strftime("%d/%m/%Y") if dados.get('data_abertura') else "N/A"
        os_recentes.append(dados)
        
    stats = {
        'total_os': total_os,
        'valor_total_gerado': valor_total_gerado,
        'os_este_mes': os_este_mes,
        'valor_este_mes': valor_este_mes,
    }
    
    return render_template(
        "dashboard.html", 
        active_page='dashboard', 
        stats=stats,
        ordens_servico=os_recentes
    )

# --- ROTAS PARA ORDEM DE SERVIÇO (OS) ---
@app.route("/ordens_servico")
@login_required
def ordens_servico_page():
    os_ref = db.collection('ordens_servico').order_by('data_abertura', direction=firestore.Query.DESCENDING).stream()
    lista_os = []
    for os in os_ref:
        dados = os.to_dict()
        dados['id'] = os.id
        dados['data_abertura_fmt'] = dados['data_abertura'].strftime("%d/%m/%Y") if dados.get('data_abertura') else "N/A"
        lista_os.append(dados)
    return render_template("ordens_servico.html", active_page='ordens_servico', lista_os=lista_os)

@app.route("/gerar_os", methods=["GET", "POST"])
@login_required
def gerar_os_page():
    if request.method == "POST":
        # A lógica de salvar a OS continua a mesma
        dados_os = {
            "cliente_nome": request.form.get('cliente_nome'),
            "cliente_doc": request.form.get('cliente_doc'),
            "cliente_endereco": request.form.get('cliente_endereco'),
            "cliente_cidade": request.form.get('cliente_cidade'),
            "cliente_telefone": request.form.get('cliente_telefone'),
            "cliente_email": request.form.get('cliente_email'),
            "servico_titulo": request.form.get('servico_titulo'),
            "servico_detalhes": request.form.get('servico_detalhes'),
            "servico_escopo": request.form.get('servico_escopo'),
            "servico_garantia": request.form.get('servico_garantia'),
            "orcamento_servicos": float(request.form.get('orcamento_servicos', 0)),
            "orcamento_produtos": float(request.form.get('orcamento_produtos', 0)),
            "data_abertura": firestore.SERVER_TIMESTAMP,
            "status": "Aberta"
        }
        db.collection('ordens_servico').add(dados_os)
        flash("Ordem de Serviço gerada com sucesso!", "success")
        return redirect(url_for('ordens_servico_page'))

    # Se for GET, busca a lista de clientes para popular o formulário
    clientes_ref = db.collection('clientes').order_by('nome').stream()
    lista_clientes = []
    for cliente in clientes_ref:
        dados = cliente.to_dict()
        dados['id'] = cliente.id
        lista_clientes.append(dados)
        
    # Converte a lista para JSON para ser usada no JavaScript
    clientes_json = json.dumps(lista_clientes)

    return render_template("gerar_os.html", active_page='gerar_os', clientes=lista_clientes, clientes_json=clientes_json)


@app.route("/visualizar_os/<os_id>")
@login_required
def visualizar_os_page(os_id):
    os_doc_ref = db.collection('ordens_servico').document(os_id)
    os_doc = os_doc_ref.get()
    if not os_doc.exists:
        return "Ordem de Serviço não encontrada!", 404
    dados_os = os_doc.to_dict()
    dados_os['id'] = os_doc.id
    dados_os['data_abertura_fmt'] = dados_os['data_abertura'].strftime("%d/%m/%Y") if dados_os.get('data_abertura') else "N/A"
    settings_ref = db.collection('configuracoes').document('empresa')
    settings_doc = settings_ref.get()
    empresa_info = settings_doc.to_dict() if settings_doc.exists else {}
    return render_template("visualizar_os.html", os=dados_os, empresa=empresa_info)

@app.route("/atualizar_status_os/<os_id>", methods=["POST"])
@login_required
def atualizar_status_os(os_id):
    novo_status = request.form.get('novo_status')
    if novo_status:
        try:
            os_ref = db.collection('ordens_servico').document(os_id)
            os_ref.update({"status": novo_status})
            flash(f"Status da O.S. atualizado para '{novo_status}'.", "success")
        except Exception as e:
            flash(f"Erro ao atualizar status: {e}", "error")
    return redirect(url_for('ordens_servico_page'))

# --- ROTAS DE CONFIGURAÇÕES ---
@app.route("/configuracoes", methods=["GET", "POST"])
@login_required
def configuracoes_page():
    settings_ref = db.collection('configuracoes').document('empresa')
    if request.method == "POST":
        form_name = request.form.get('form_name')
        if form_name == 'change_password':
            new_password, confirm_password = request.form.get('new_password'), request.form.get('confirm_password')
            if not new_password or new_password != confirm_password:
                flash("As senhas não coincidem ou estão vazias.", "error")
            else:
                rest_api_url = f"https://identitytoolkit.googleapis.com/v1/accounts:update?key={FIREBASE_WEB_API_KEY}"
                payload = json.dumps({"idToken": session.get('idToken'), "password": new_password, "returnSecureToken": True})
                try:
                    r = requests.post(rest_api_url, data=payload)
                    r.raise_for_status()
                    session['idToken'] = r.json().get('idToken')
                    flash("Senha alterada com sucesso!", "success")
                except Exception:
                    flash("Sua sessão pode ter expirado. Faça login novamente para alterar a senha.", "error")
        elif form_name == 'company_info':
            company_data = {'cnpj': request.form.get('cnpj'), 'endereco': request.form.get('endereco'), 'nome_empresa': request.form.get('nome_empresa')}
            settings_ref.set(company_data)
            flash("Informações da empresa salvas com sucesso!", "success")
        return redirect(url_for('configuracoes_page'))
    
    current_settings = (settings_ref.get()).to_dict() if settings_ref.get().exists else {}
    return render_template("configuracoes.html", active_page='configuracoes', settings=current_settings)

# --- ROTAS ANTIGAS DE CONTRATO (Manter se desejar) ---
@app.route("/gerar", methods=["GET", "POST"])
@login_required
def gerar_page():
    # ... (código da antiga rota de gerar contrato) ...
    flash("Esta função foi substituída pelas Ordens de Serviço.", "error")
    return redirect(url_for('gerar_os_page'))

@app.route("/contratos")
@login_required
def contratos_page():
    # ... (código da antiga rota de listar contratos) ...
     flash("Esta função foi substituída pelas Ordens de Serviço.", "error")
     return redirect(url_for('ordens_servico_page'))

# Adicione este bloco de código em app.py

# ================================================================
# --- ROTAS PARA CARTEIRA DE CLIENTES ---
# ================================================================

@app.route("/clientes")
@login_required
def clientes_page():
    # Busca todos os clientes, ordenados por nome
    clientes_ref = db.collection('clientes').order_by('nome').stream()
    lista_clientes = []
    for cliente in clientes_ref:
        dados = cliente.to_dict()
        dados['id'] = cliente.id
        lista_clientes.append(dados)
    return render_template("clientes.html", active_page='clientes', clientes=lista_clientes)

@app.route("/adicionar_cliente", methods=["POST"])
@login_required
def adicionar_cliente():
    try:
        dados_cliente = {
            "nome": request.form.get('nome'),
            "doc": request.form.get('doc'),
            "endereco": request.form.get('endereco'),
            "cidade": request.form.get('cidade'),
            "telefone": request.form.get('telefone'),
            "email": request.form.get('email'),
        }
        db.collection('clientes').add(dados_cliente)
        flash("Cliente adicionado com sucesso!", "success")
    except Exception as e:
        flash(f"Erro ao adicionar cliente: {e}", "error")
    return redirect(url_for('clientes_page'))

@app.route("/editar_cliente/<cliente_id>", methods=["POST"])
@login_required
def editar_cliente(cliente_id):
    try:
        dados_atualizados = {
            "nome": request.form.get('nome'),
            "doc": request.form.get('doc'),
            "endereco": request.form.get('endereco'),
            "cidade": request.form.get('cidade'),
            "telefone": request.form.get('telefone'),
            "email": request.form.get('email'),
        }
        db.collection('clientes').document(cliente_id).update(dados_atualizados)
        flash("Cliente atualizado com sucesso!", "success")
    except Exception as e:
        flash(f"Erro ao atualizar cliente: {e}", "error")
    return redirect(url_for('clientes_page'))

@app.route("/excluir_cliente/<cliente_id>", methods=["POST"])
@login_required
def excluir_cliente(cliente_id):
    try:
        db.collection('clientes').document(cliente_id).delete()
        flash("Cliente excluído com sucesso!", "success")
    except Exception as e:
        flash(f"Erro ao excluir cliente: {e}", "error")
    return redirect(url_for('clientes_page'))


# --- INICIALIZAÇÃO DO SERVIDOR ---
if __name__ == "__main__":
    app.run(debug=True)
